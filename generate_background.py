from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Self-Checking Summarizer for Long Government Documents')
run.bold = True
run.font.size = Pt(14)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CS6140 Machine Learning — Final Project Report')
run.italic = True
run.font.size = Pt(12)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.add_run('Author: ______________________')

doc.add_paragraph()

h1 = doc.add_heading('2. Background', level=1)

# 2.1
doc.add_heading('2.1 From Chatbot to Agent Workflow', level=2)

p = doc.add_paragraph()
p.add_run(
    "Large language models (LLMs) first entered broad public use as conversational "
    "chatbots. In this paradigm, the model consumes a prompt and emits a single "
    "response in one forward pass; any reasoning, retrieval, or multi-step "
    "composition must be encoded inside that single exchange. The design is simple "
    "and responsive, but it struggles whenever the task exceeds what can be solved "
    "in one read-through — long documents, multi-hop reasoning, or problems that "
    "benefit from iterative refinement."
)

p = doc.add_paragraph()
p.add_run(
    "The community response has been a shift toward "
)
r = p.add_run("agent workflows")
r.italic = True
p.add_run(
    ". An agentic system treats the LLM not as a single oracle but as a reusable "
    "computational unit that can be orchestrated through explicit control flow: "
    "the model is called multiple times with different prompts, produces intermediate "
    "artifacts that feed downstream calls, and can be paired with tools, verifiers, "
    "and planners. Representative building blocks include chain-of-thought prompting, "
    "tool use and function calling, ReAct-style planner–executor loops, and "
    "map-reduce or refine pipelines over long inputs. The common thread is that a "
    "hard task is decomposed into smaller, individually tractable subtasks, each of "
    "which fits comfortably inside one model call."
)

p = doc.add_paragraph()
p.add_run(
    "For long-document summarization — the setting of this project — the agentic "
    "view is especially natural. A 15k-token government report cannot be both fully "
    "read and densely compressed in a single pass, even with models whose context "
    "windows nominally cover the input: in practice, the output still behaves as "
    "though the model only has bandwidth for the earliest portion of the document. "
    "Treating summarization as an agent workflow — segment, summarize locally, "
    "integrate globally, then verify — moves the problem back inside the model's "
    "effective capacity."
)

# 2.2
doc.add_heading('2.2 Previous Work: One-Shot Inference Baseline', level=2)

p = doc.add_paragraph()
p.add_run(
    "Before designing the divide-and-conquer pipeline, we first verified that the "
    "chosen backbone — Qwen/Qwen2.5-7B-Instruct under 4-bit NF4 quantization — was "
    "usable on the commodity hardware available to the project (a Tesla T4 with "
    "roughly 15 GB of VRAM). This verification step is documented in "
)
r = p.add_run("CS6140_inference_verification.ipynb")
r.font.name = 'Consolas'
p.add_run(
    ". The notebook loads the model through the Hugging Face Transformers stack "
    "with BitsAndBytes 4-bit quantization, confirms that the chat template can "
    "generate coherent long-form technical explanations (for example, a multi-"
    "paragraph response on linear regression, with properly formatted equations "
    "and itemized modeling assumptions), and then frees GPU memory cleanly."
)

p = doc.add_paragraph()
p.add_run(
    "This established two facts that the rest of the project depends on. First, "
    "the backbone is fluent and instruction-following: any subsequent failures on "
    "long-document summarization cannot be blamed on an inability to write good "
    "English or to follow structured prompts. Second, the backbone is affordable: "
    "a single inference pass fits on one T4, which makes repeated multi-stage "
    "agent calls computationally feasible within Colab and Northeastern HPC time "
    "limits."
)

p = doc.add_paragraph()
p.add_run(
    "The broader literature reinforces the same direction. Stepwise, retrieval-"
    "augmented, and map-reduce style summarization pipelines have repeatedly been "
    "shown to outperform naive truncation on long inputs such as GovReport, arXiv, "
    "and PubMed; and self-checking or \"LLM-as-a-judge\" patterns have emerged as "
    "a cheap proxy for human faithfulness evaluation when annotation budgets are "
    "limited. Our work sits squarely inside this line of research, combining a "
    "decomposition-based summarizer with a model-as-critic verification stage."
)

# 2.3
doc.add_heading('2.3 What Makes This Work Particularly Interesting', level=2)

p = doc.add_paragraph()
p.add_run(
    "Two design choices distinguish this project from a generic \"long "
    "summarization with a big model\" study."
)

p = doc.add_paragraph()
p.add_run("First, the pipeline is explicitly framed as an application of the classical ")
r = p.add_run("divide-and-conquer")
r.bold = True
p.add_run(
    " algorithmic strategy — the same principle behind merge sort, the Fast Fourier "
    "Transform, and many dynamic-programming recurrences. A long report is split "
    "into smaller sub-problems (chunks); each sub-problem is solved independently "
    "by the same model (the map step); and the partial results are then combined "
    "by a dedicated reducer into a single coherent summary. The analogy is not "
    "only cosmetic: the recursive-then-combine structure is precisely what lets a "
    "fixed-capacity model handle inputs of arbitrary length, in the same way that "
    "merge sort lets a fixed comparator handle arbitrarily long arrays. Framing "
    "the summarizer this way yields a clean, theoretically motivated scaffold — "
    "segmenter, mapper, reducer, and optional critic/refiner — against which "
    "individual design choices can be ablated."
)

p = doc.add_paragraph()
p.add_run("Second, the choice of backbone is deliberately ")
r = p.add_run("not")
r.italic = True
p.add_run(
    " a frontier model. Qwen2.5-7B-Instruct is a mid-sized, 4-bit-quantized open "
    "model. It is not expected to compete on raw capability with GPT-4-class "
    "systems, and any single-call baseline it produces will inevitably under-cover "
    "a 15k-token document. "
)
r = p.add_run("That is the point.")
r.bold = True
p.add_run(
    " By fixing the backbone at a level where the one-shot baseline is visibly "
    "imperfect, the project can isolate and measure the contribution of the "
)
r = p.add_run("method")
r.italic = True
p.add_run(
    " — the segmentation, the map-reduce structure, the self-check, the refiner — "
    "rather than riding on the raw competence of a stronger LLM. If coverage, "
    "faithfulness, or ROUGE improve under the same backbone, the gain is "
    "attributable to the pipeline, not to the model. This turns the study from a "
    "capability demonstration into a controlled experiment on agent-workflow "
    "design, which is the more transferable contribution of the work."
)

# 2.4
doc.add_heading('2.4 Dataset: GovReport', level=2)

p = doc.add_paragraph()
p.add_run(
    "The experiments in this project are conducted on "
)
r = p.add_run("GovReport")
r.italic = True
p.add_run(
    ", a long-document summarization benchmark built from publicly released "
    "reports of the U.S. Government Accountability Office (GAO) and the "
    "Congressional Research Service (CRS). Each instance pairs a full report "
    "(routinely exceeding 10,000 tokens) with a human-written executive summary "
    "authored by the same agency that produced the report. Two properties of "
    "this dataset make it an especially appropriate testbed for the divide-and-"
    "conquer pipeline proposed in this work."
)

p = doc.add_paragraph()
p.add_run("First, the corpus is ")
r = p.add_run("textually clean and structurally regular")
r.bold = True
p.add_run(
    ". Unlike web-scraped news, social media, or transcribed speech, GAO and CRS "
    "documents are professionally edited publications written in a formal "
    "bureaucratic register. They follow consistent section conventions — "
    "background, findings, evidence, recommendations, and agency comments — and "
    "are largely free of the boilerplate, advertising artifacts, OCR errors, "
    "colloquialisms, and conversational noise that typically contaminate "
    "open-domain summarization corpora. For a study whose goal is to isolate the "
    "contribution of pipeline design, this cleanliness is a methodological asset: "
    "it suppresses a whole class of confounding failure modes (broken encodings, "
    "off-topic digressions, stylistic inconsistency) so that differences in "
    "output quality can be more confidently attributed to the segmentation, "
    "mapping, and reduction stages rather than to upstream noise in the input."
)

p = doc.add_paragraph()
p.add_run("Second, the ")
r = p.add_run("logical structure of a government report is complex enough")
r.bold = True
p.add_run(
    " that faithful summarization cannot be shortcut by the model's parametric "
    "world knowledge. Reports routinely interleave statute-specific definitions, "
    "agency-specific acronyms, tabulated statistics, multi-step causal chains "
    "between policy actions and observed outcomes, and qualified "
    "recommendations whose scope depends on conditions introduced dozens of "
    "pages earlier. A summary that is faithful to one of these documents must "
    "track entity references across long spans, respect quantitative claims, "
    "and preserve the argumentative order linking findings to recommendations — "
    "none of which the LLM can fabricate plausibly from prior training. In "
    "other words, the dataset punishes any strategy that collapses into \"write "
    "something that sounds like a GAO summary\" and rewards strategies that "
    "actually read and reconcile the whole document. This is precisely the "
    "behavior that a divide-and-conquer pipeline, with its explicit map-reduce "
    "and self-check stages, is designed to exhibit, which is why GovReport is "
    "a more discriminative benchmark for the method than shorter or more "
    "formulaic summarization corpora such as CNN/DailyMail or XSum."
)

# =============================================================
# 3. Methodology
# =============================================================

def add_pseudocode(doc, lines):
    """Render a block of pseudocode lines in a monospace font, preserving indentation."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line if line else " ")  # non-breaking space for blank lines
        r.font.name = 'Consolas'
        r.font.size = Pt(10)

doc.add_heading('3. Methodology', level=1)

# 3.1 Overview
doc.add_heading('3.1 Pipeline Overview', level=2)

p = doc.add_paragraph()
p.add_run(
    "Our summarization system is a four-stage divide-and-conquer agent built on "
    "top of a single frozen backbone (Qwen2.5-7B-Instruct, 4-bit NF4). Each stage "
    "is a distinct prompt template issued to the same model; no fine-tuning is "
    "performed, so all behavioral differences across stages come from prompt "
    "design and control flow rather than from model weights. The four stages are "
    "(i) "
)
r = p.add_run("chunking"); r.italic = True
p.add_run(", which splits the full source report into overlapping token windows; (ii) ")
r = p.add_run("map"); r.italic = True
p.add_run(", which asks the model to summarize each chunk independently with positional context; (iii) ")
r = p.add_run("reduce"); r.italic = True
p.add_run(
    ", which merges the partial summaries into a single coherent output, falling "
    "back to a hierarchical merge when the concatenated partials exceed the "
    "reducer's context budget; and (iv) a "
)
r = p.add_run("refine variant"); r.italic = True
p.add_run(
    " that replaces the map-plus-reduce structure with a running summary that is "
    "iteratively updated as the model walks through the chunks in order. At "
    "inference time, the pipeline can be run in either MapReduce mode or Refine "
    "mode; the two share the chunking stage and the backbone but differ in how "
    "chunk-level information is aggregated."
)

# 3.2 Chunking
doc.add_heading('3.2 Chunking', level=2)

p = doc.add_paragraph()
p.add_run(
    "Unlike the one-shot baseline, which truncates the input to the first 16k "
    "tokens and therefore discards the tail of most GovReport documents, the "
    "divide-and-conquer pipeline processes every token of the source. Each "
    "report is tokenized with the Qwen tokenizer and split into a sequence of "
    "fixed-size windows of "
)
r = p.add_run("3,000 tokens"); r.bold = True
p.add_run(" each, with ")
r = p.add_run("300 tokens of overlap"); r.bold = True
p.add_run(
    " (roughly one paragraph) between consecutive windows. The chunk size is "
    "deliberately set well below the model's context limit so that the entire "
    "chunk, together with the map prompt and the generated chunk summary, fits "
    "inside a single forward pass without displacing instruction tokens. The "
    "overlap is a guard against information loss at chunk boundaries: a "
    "sentence that straddles two windows is guaranteed to appear in at least "
    "one of them in full, which is important for multi-sentence findings and "
    "causal claims whose meaning cannot be recovered from a truncated half."
)

doc.add_paragraph().add_run("Pseudocode.").italic = True
add_pseudocode(doc, [
    "function CHUNK_REPORT(report, chunk_size = 3000, overlap = 300):",
    "    tokens  = tokenize(report)",
    "    chunks  = []",
    "    start   = 0",
    "    while start < len(tokens):",
    "        end   = min(start + chunk_size, len(tokens))",
    "        chunks.append(detokenize(tokens[start:end]))",
    "        if end == len(tokens): break",
    "        start = start + (chunk_size - overlap)",
    "    return chunks",
])

# 3.3 Map
doc.add_heading('3.3 Map Phase — Local Chunk Summarization', level=2)

p = doc.add_paragraph()
p.add_run(
    "In the map phase, each chunk is summarized independently by the backbone. "
    "The prompt is constructed so that the model is explicitly told it is "
    "looking at part "
)
r = p.add_run("i"); r.italic = True
p.add_run(" of ")
r = p.add_run("N"); r.italic = True
p.add_run(
    "; this positional framing discourages the model from hallucinating global "
    "conclusions (\"the report concludes that…\") from what is in fact only a "
    "local fragment, and instead steers it toward extracting the findings, "
    "quantitative evidence, and policy recommendations that are actually "
    "present in the chunk. Decoding uses a low temperature (T = 0.2) to "
    "prioritize factual extraction over stylistic variation, and the output "
    "budget is capped at 300 new tokens per chunk so that the total map cost "
    "scales linearly in the number of chunks rather than in the document length "
    "squared. All chunks are processed sequentially by the same model instance, "
    "with no information flowing between chunks — this independence is what "
    "distinguishes map from refine."
)

doc.add_paragraph().add_run("Pseudocode.").italic = True
add_pseudocode(doc, [
    "function MAP_PHASE(report):",
    "    chunks           = CHUNK_REPORT(report)",
    "    partial_summaries = []",
    "    for i in 0 .. len(chunks) - 1:",
    "        prompt = \"Summarize part {i+1} of {N} of a government report. \"",
    "                 \"Extract findings, data, and recommendations from this section.\"",
    "                 + chunks[i]",
    "        partial_summaries.append( LLM(prompt, T = 0.2, max_new = 300) )",
    "    return chunks, partial_summaries",
])

# 3.4 Reduce
doc.add_heading('3.4 Reduce Phase — Global Merge (with Hierarchical Fallback)', level=2)

p = doc.add_paragraph()
p.add_run(
    "The reduce phase consumes the list of partial summaries produced by the "
    "map phase and emits a single coherent summary. Partials are concatenated "
    "with explicit section markers (\"[Section i]\") and passed to the model "
    "with an instruction that asks it to (a) eliminate redundancy introduced by "
    "the chunk overlap, (b) preserve all key findings, quantitative evidence, "
    "and policy recommendations, and (c) hit a target length of 400–600 words, "
    "which is close to the empirical reference length in GovReport. A slightly "
    "higher temperature (T = 0.3) is used here than in the map phase, because "
    "the reducer's job includes paraphrase and integration rather than pure "
    "extraction."
)

p = doc.add_paragraph()
p.add_run("When the concatenated partials exceed an 8,000-token context budget — a regime that arises for the longest GovReport documents where a single reduce call would otherwise overflow — the pipeline falls back to a ")
r = p.add_run("hierarchical reduce"); r.bold = True
p.add_run(
    ". The partials are grouped into triples, each triple is merged into an "
    "intermediate summary by a smaller-budget prompt, and the resulting "
    "intermediates are themselves fed back into the reduce routine. This recurses "
    "until the surviving set of summaries fits within the context budget, at "
    "which point a single final reduce call produces the output summary. "
    "Conceptually, this is the same binary-tree reduction that classical "
    "map-reduce systems apply when a single reducer cannot hold the full "
    "intermediate set in memory."
)

doc.add_paragraph().add_run("Pseudocode.").italic = True
add_pseudocode(doc, [
    "function REDUCE(partials, context_limit = 8000, max_new = 800):",
    "    combined = join(partials, separator = \"[Section i]\")",
    "    if token_length(combined) <= context_limit:",
    "        prompt = \"Synthesize these section summaries into a single coherent summary. \"",
    "                 \"Remove redundancy; preserve findings, data, recommendations. \"",
    "                 \"Target length: 400-600 words.\" + combined",
    "        return LLM(prompt, T = 0.3, max_new = max_new)",
    "    else:",
    "        return HIERARCHICAL_REDUCE(partials, context_limit, max_new)",
    "",
    "function HIERARCHICAL_REDUCE(partials, context_limit, max_new):",
    "    current = partials",
    "    repeat:",
    "        groups    = partition(current, size = 3)",
    "        next_level = []",
    "        for g in groups:",
    "            if token_length(join(g)) <= context_limit:",
    "                next_level.append( LLM(\"Merge into one concise summary.\" + join(g),",
    "                                       T = 0.2, max_new = 400) )",
    "            else:",
    "                next_level.extend(g)        # pass through oversized items",
    "        if token_length(join(next_level)) <= context_limit or len(next_level) == 1:",
    "            return REDUCE(next_level, context_limit, max_new)",
    "        current = next_level",
])

# 3.5 Refine
doc.add_heading('3.5 Refine Variant — Iterative Running Summary', level=2)

p = doc.add_paragraph()
p.add_run(
    "The refine variant replaces the independent-then-merge structure of "
    "MapReduce with a single sequential pass that maintains a "
)
r = p.add_run("running summary"); r.italic = True
p.add_run(
    ". The first chunk is summarized with a standard section-summary prompt, "
    "producing an initial running summary. For every subsequent chunk, the "
    "prompt supplies the current running summary alongside the next chunk and "
    "asks the model to update the summary by incorporating new information "
    "from that chunk while keeping the output within the same 400–600 word "
    "target. After the final chunk, the running summary is returned as the "
    "model's output. This design has two distinguishing properties. First, it "
    "eliminates the explicit reduce step, because integration happens "
    "continuously as the model reads the document, which removes one potential "
    "source of paraphrase-induced precision loss. Second, the influence of "
    "each chunk on the final output is asymmetric: later chunks can rewrite or "
    "delete content contributed by earlier chunks, but not vice versa, which "
    "trades the symmetry of MapReduce for the tighter inter-section coherence "
    "of a sequential read."
)

doc.add_paragraph().add_run("Pseudocode.").italic = True
add_pseudocode(doc, [
    "function REFINE_SUMMARIZE(report):",
    "    chunks          = CHUNK_REPORT(report)",
    "    running_summary = \"\"",
    "    for i in 0 .. len(chunks) - 1:",
    "        if i == 0:",
    "            prompt = \"Summarize this section. Focus on findings and recommendations.\"",
    "                     + chunks[i]",
    "        else:",
    "            prompt = \"Below is the running summary so far, followed by the next section.\"",
    "                     \"Update the summary to incorporate new information. \"",
    "                     \"Keep target length 400-600 words.\"",
    "                     + \"CURRENT SUMMARY:\" + running_summary",
    "                     + \"NEXT SECTION:\"    + chunks[i]",
    "        running_summary = LLM(prompt, T = 0.3, max_new = 700)",
    "    return running_summary",
])

# Wrap-up paragraph
p = doc.add_paragraph()
p.add_run(
    "Both pipelines are run over the same test slice of GovReport with the same "
    "backbone, the same tokenizer, the same chunk-size and overlap settings, "
    "and the same self-check verifier described in §2; they differ only in "
    "whether chunk-level information is aggregated in parallel (MapReduce with "
    "its explicit reduce stage) or sequentially (Refine with its running "
    "summary). This shared scaffolding is what lets the two variants serve as "
    "a controlled ablation of how the aggregation strategy, as opposed to the "
    "backbone or the chunking, affects final summary quality."
)

out_path = r"c:\Users\usp78\Desktop\NEU_docs\CS6140\Self_checking_summarizer\CS6140_report_WIP.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
